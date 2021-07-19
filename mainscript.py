# -*- coding: utf-8 -*-

import eel
from random import choice
import docx
from docx import Document
import os
import pdfplumber
import tkinter as tk
from tkinter import filedialog



intdocx = 1
myqs = docx.Document()
qstring = ""
zaminefiles= ""
adr = ""
boolforcheckpdfword = False
pdf_word_input = ""

#----------------- pdf reader -----------------------
# pdfread = ""
# with pdfplumber.open("1.pdf") as pdf:
#     first_page = pdf.pages[0]
#     pdfread += first_page.extract_text()

# print(pdfread)
#----------------- pdf reader -----------------------

#----------------- word reader -----------------------

#     try:
#         doc = docx.Document(zaminefiles)  # Creating word reader object.
#         data = ""
#         fullText = []
#         for para in doc.paragraphs:
#             fullText.append(para.text)
#             data = '\n'.join(fullText)
 
#         print(data)
 
#     except IOError:
#         print('There was an error opening the file!')
#         return
        
#----------------- word reader -----------------------


@eel.expose

def gettext(nextqs,numqst,numqsf,namef):
    global qstring
    global zaminefiles
    global adr
    global boolforcheckpdfword
    global pdf_word_input


    splitmetod = "<__>nextq<__>"




    if (boolforcheckpdfword):
        splitmetod = "<next>"
        if(".pdf" in pdf_word_input):
            #----------------- pdf reader -----------------------
            pdfread = ""
            with pdfplumber.open(pdf_word_input) as pdf:
                first_page = pdf.pages[0]
                pdfread += first_page.extract_text()

            print(pdfread)
            #----------------- pdf reader -----------------------
            nextqs = pdfread
            
        if(".docx" in pdf_word_input):
            #----------------- word reader -----------------------

                try:
                    doc = docx.Document(pdf_word_input)  # Creating word reader object.
                    data = ""
                    fullText = []
                    for para in doc.paragraphs:
                        fullText.append(para.text)
                        data = '\n'.join(fullText)
 
                    print(data)
                    nextqs = str(data)

                except IOError:
                    print('There was an error opening the file!')
                    return
        
            #----------------- word reader -----------------------

    else:
        splitmetod = "<__>nextq<__>"


    try:
        doc = docx.Document(zaminefiles)  # Creating word reader object.
        data = ""
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
            data = '\n'.join(fullText)
 
        print(data)
 
    except IOError:
        print('There was an error opening the file!')
        return
        
    qsnum = int(numqst)-1
    lisnumwhile = -1
    files = 1
    qs = str(nextqs).split(splitmetod)
    if(boolforcheckpdfword):
        print("")
    else:
        qs.remove("")
    numf = int(numqsf)
    numf = numf+1
    namefs = namef



    while  files < numf :
        
        myqs = docx.Document()
        myqs.add_paragraph(data+'\n')
        
        while qsnum > lisnumwhile :

            lisnumwhile = lisnumwhile+1
            print(qs)

            dlist = choice(qs)

            qs.remove(dlist)
            print(dlist)

            if("<~" in dlist):        
                imgs = dlist.split("<~")
                for imgsplit in imgs :
                    print(imgs)
                    if(".jpg" in imgsplit or ".png" in imgsplit):
                        print("img")
                        myqs.add_picture(imgsplit)
                    else:
                        print("________++++"+imgsplit)
                        myqs.add_paragraph(imgsplit)
                    
            else:
                myqs.add_paragraph(dlist)
            
            
            
        print("------------------------------")

        newpath = str(adr)+"\_"+str(namefs) 
        if not os.path.exists(newpath):
            os.mkdir(newpath)

        
        myqs.save(str(adr)+"/"+"_"+str(namefs)+"/" +str(namefs)+str(files) +".docx")
        files = files+1
        lisnumwhile = -1
        myqs = 0
        qs = str(nextqs).split(splitmetod)
        if(boolforcheckpdfword):
            print("")
        else:
            qs.remove("")


@eel.expose
def zaminefile():
    global zaminefiles
    #---------------------- file dialog ---------------------
    root = tk.Tk()
    root.withdraw()
    zaminefiles = filedialog.askopenfilename(title="فایل زمینه",
    filetypes=[('فایل زمینه رو انتخاب کن', '.docx')])
    #---------------------- file dialog ---------------------



@eel.expose
def filesavefolder():
    global adr
    root = tk.Tk()
    root.withdraw()
    savefolder = filedialog.askdirectory()
    adr = savefolder


@eel.expose
def checkboxwordpdfinput(checkboxwordpdfinput):
    global boolforcheckpdfword
    if(checkboxwordpdfinput == False):
        boolforcheckpdfword = False  
    if(checkboxwordpdfinput== True):
        boolforcheckpdfword = True  


@eel.expose
def pdf_word_input():
    global pdf_word_input
    root = tk.Tk()
    root.withdraw()
    pdf_word_input = filedialog.askopenfilename(title="ورودی سوالات از طریق فایل",
    filetypes=[('سوالات word فایل ', '.docx'),('سوالات pdf فایل ', '.pdf')])


eel.init('web')
eel.start('main_page.html' , size=(700,500) )  # Start