from os import mkdir, name, read, times
import os

import docx
from docx.api import Document
import random 
from random import choice
from docx import Document
import time
import os
from shutil import copyfile

class prossesqs:


    def output_ps(filesrc,numberfile,numberqs,name,sarbarg,savesrc):

        try:
            os.mkdir(f"{savesrc}\\{name}")
        except:
            pass

        textsq = ""
        with open(filesrc , "r+", encoding="utf-8") as sqfile:
            textsq = sqfile.read()

        textsq = textsq.split("{/_-next-qs-_`}")
        textsq.remove('')


        try:
            for filemk in range(0,int(numberfile)):
                myqs = docx.Document()

                if len(sarbarg) > 0 :
                    for e in Document(sarbarg).element.body:
                        myqs.element.body.append(e)

                myqs.save(f"{savesrc}\\{name}\\درحال پردازش..{filemk}.docx")
                
                myqs = docx.Document()

                if len(textsq) >= int(numberqs):
                    for sqrandom in range(0,int(numberqs)):

                        randomselect = choice(textsq)

                        if("<-- img_imported_cod -->" in randomselect):        
                            imgs = randomselect.split("<-- img_imported_cod -->")
                            for imgsplit in imgs :
                                if(".jpg" in imgsplit or ".png" in imgsplit):
                                    myqs.add_picture(imgsplit)
                                else:
                                    myqs.add_paragraph(imgsplit)
                                
                        else:
                            myqs.add_paragraph(str(randomselect+"\n"))
                        textsq.remove(randomselect)

                myqs.save(f"{savesrc}\\{name}\\درحال پردازش...{filemk}.docx")

                time.sleep(0.1)
                myqs = docx.Document()

                for e in Document(f"{savesrc}\\{name}\\درحال پردازش..{filemk}.docx").element.body:
                    myqs.element.body.append(e)
                for e in Document(f"{savesrc}\\{name}\\درحال پردازش...{filemk}.docx").element.body:
                    myqs.element.body.append(e)

                myqs.save(f"{savesrc}\\{name}\\{name}{filemk}.docx")

                os.remove(f"{savesrc}\\{name}\\درحال پردازش..{filemk}.docx")
                os.remove(f"{savesrc}\\{name}\\درحال پردازش...{filemk}.docx")

                textsq = ""
                with open(filesrc , "r+", encoding="utf-8") as sqfile:
                    textsq = sqfile.read()

                textsq = textsq.split("{/_-next-qs-_`}")
                textsq.remove('')
        except:
            pass
            print(".. error in Export ..")

    def import_sovalsazfile(src):
        if src != "":
            namefile = src.split("/")
            namefile = namefile[len(namefile)-1]
            copyfile(src, f"QS\\{namefile}")


    def wordORpdf_to_qsapp():
        pass
